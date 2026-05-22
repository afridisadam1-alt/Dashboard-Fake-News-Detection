# =========================================================
# app.py – Disinformation Detection Dashboard (DL + ML Cached, Toolbar Hidden)
# =========================================================

import os, warnings, pickle, re
import streamlit as st
import pandas as pd
import numpy as np
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from tensorflow.keras.models import load_model
from tensorflow.keras.preprocessing.sequence import pad_sequences
import altair as alt
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import tensorflow as tf
import hashlib
def set_active_source(source, text):
    st.session_state.active_source = source
    st.session_state.input_text = clean_text(text)
    st.session_state.text_hash = None
    st.session_state.summary = None
    st.session_state.prediction_done = False
if "text_hash" not in st.session_state:
    st.session_state.text_hash = None


def get_hash(text):
    return hashlib.md5(text.encode()).hexdigest()
# =========================
# SESSION STATE INIT (CLEAN)
# =========================
defaults = {
    "top_keywords": "",
    "summary": None,
    "input_text": "",
    "prediction_done": False,
    "input_source": "dataset",
    "active_source": "dataset",
    "current_label": None,
    "last_zip_file": None,
     "zip_initialized": False,
    "zip_user_selected": False,
    "questionnaire_submitted": False
}

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

#Gemini API
import google.generativeai as genai
def set_source(source):
    st.session_state.input_source = source
    st.session_state.active_source = source   # 🔥 IMPORTANT FIX
    st.session_state.input_text = ""
    st.session_state.prediction_done = False
    st.session_state.summary = None
    st.session_state.current_label = None

try:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    gemini_model = genai.GenerativeModel("models/gemini-2.5-flash")
except Exception:
    st.error("Gemini API not available. Please check API key.")
    gemini_model = None
# =========================================================
# SYSTEM INTRO (HERO SECTION)
# =========================================================
st.markdown("""
<style>
.big-title {
    font-size: 42px;
    font-weight: 800;
    color: #1f4e79;
    text-align: center;
    margin-bottom: 10px;
}

.sub-text {
    font-size: 18px;
    text-align: center;
    color: #444;
    margin-bottom: 25px;
    line-height: 1.6;
}

.author-box {
    font-size: 16px;
    text-align: center;
    color: #666;
    margin-top: 10px;
}
            /* Expander title (this is the clickable header) */
[data-testid="stExpander"] summary {
    font-size: 22px !important;
    font-weight: 700;
    color: #1f4e79;
}
</style>

<div class="big-title">
🧠 AI-Based Disinformation Detection System
</div>

<div class="sub-text">
A research-driven platform combining Machine Learning (ML) and Deep Learning (BiLSTM) 
to detect whether online news content is <b>True</b> or <b>Disinformation</b>.<br><br>

The system integrates text preprocessing, predictive modeling, and explainable AI visualizations 
to support transparent and interpretable decision-making.
</div>

<div class="author-box">
<b>Developed by:</b> Mr Sadam Hussain (PhD Sholar) | Disinformation Detection & AI Systems<br>
<b>Purpose:</b> Academic research on fake news detection, user trust, and human-AI interaction
</div>
            
""", unsafe_allow_html=True)
st.markdown("""
## 📘 User Testing Instructions

1. Select ML or DL model (top left sidebar)
2. Search dataset content using the search box or radio buttons below
3. Predict results using one of the following:

    - a) Select a record from the dataset (radio buttons)
    - b) Select from ZIP folder (preloaded news articles)
    - c) Upload your own file (TXT, PDF, DOCX, CSV)

4. Click **Predict**
5. Review explanation visualizations
6. Understand the prediction result with AI-generated summary
7. Complete the feedback questionnaire
""")
# =========================================================
# Check package versions
# =========================================================
print("streamlit:", st.__version__)
print("pandas:", pd.__version__)
print("numpy:", np.__version__)
print("PyPDF2:", PdfReader.__module__.split('.')[0], "version not directly accessible")
print("python-docx:", Document.__module__.split('.')[0], "version not directly accessible")
print("tensorflow:", tf.__version__)
print("altair:", alt.__version__)
print("wordcloud:", WordCloud.__module__.split('.')[0], "version not directly accessible")

# =========================================================
# System cleanup
# =========================================================
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "3"
warnings.filterwarnings("ignore")
st.set_page_config(page_title="Disinformation Detection Dashboard", layout="wide")

# =========================================================
# Constants
# =========================================================
MAX_LEN = 700

# =========================================================
# Text cleaning function
# =========================================================
def clean_text(text):
    """
    Removes special/garbled characters and extra whitespace
    """
    text = re.sub(r"[^\x00-\x7F]+", " ", text)  # remove non-ASCII
    text = re.sub(r"\s+", " ", text)            # collapse multiple spaces
    return text.strip()

# =========================================================
# Dataset configs (local files only)
# =========================================================
ml_datasets = {
"EUvsIGF": {"model": "euvsipf_pac.pkl", "vectorizer": "euvsipf_vectorizer.pkl", "csv": "EUvsIPF-Dataset_sample20.csv", "text_col": "text_english"},
}

dl_datasets = {

    "EUvsIGF": {"model": "EUvsIPF_bilstm.h5", "tokenizer": "EUvsIPF_tokenizer.pkl", "csv": "EUvsIPF-Dataset_sample20.csv", "text_col": "text_english"},
}

# =========================================================
# Prediction label mapping
# =========================================================
PRED_LABEL_MAP = {
    "FA-KES": {0: "TRUE", 1: "Fake"},
    "ISOT": {0: "True", 1: "Fake"},
    "EUvsISOT": {0: "True", 1: "disinformation"},
    "George McIntire": {0: "REAL", 1: "FAKE"},
    "EUvsIGF": {0: "true", 1: "disinformation"},
    "EUvsDisinfo": {0: "support", 1: "disinformation"},
}

DL_POSITIVE_LABEL = {
    "FA-KES": PRED_LABEL_MAP["FA-KES"],
    "ISOT": PRED_LABEL_MAP["ISOT"],
    "EUvsISOT": PRED_LABEL_MAP["EUvsISOT"],
}

# =========================================================
# Label normalization
# =========================================================
def normalize_prediction(dataset_name, pred_label):
    map_ = PRED_LABEL_MAP.get(dataset_name, {})
    if isinstance(pred_label, str) and pred_label.isdigit():
        pred_label = int(pred_label)
    mapped = map_.get(pred_label, pred_label)
    mapped_lower = str(mapped).strip().lower()
    if mapped_lower in ["true", "real", "support", "0"]:
        return "True"
    elif mapped_lower in ["fake", "1"]:
        return "Fake"
    elif mapped_lower in ["disinformation"]:
        return "Disinformation"
    return str(mapped)

# =========================================================
# Detect label column
# =========================================================
def detect_label_column(df, dataset_name):
    if dataset_name == "FA-KES": return "labels"
    elif dataset_name == "ISOT": return "label"
    elif dataset_name == "EUvsISOT": return "class"
    else:
        for col in ["class", "label", "labels"]:
            if col in df.columns: return col
    return None

# =========================================================
# Load dataset (local only)
# =========================================================
@st.cache_data(show_spinner=False)
def load_dataset(cfg, dataset_name):
    csv_file = cfg["csv"]
    df = pd.read_csv(csv_file)
    label_col = detect_label_column(df, dataset_name)
    text_col = cfg["text_col"]
    df[label_col] = pd.Series(df[label_col], dtype=str).str.strip()
    df[text_col] = pd.Series(df[text_col], dtype=str).str.strip()
    df[text_col] = df[text_col].apply(clean_text)  # <-- clean dataset text
    df = df[~df[label_col].isin(["", "nan", "NaN"])]
    df = df[df[text_col] != ""]
    df = df[[text_col, label_col]].copy()
    df["Select"] = False
    return df, label_col

# =========================================================
# Load ML model (local only)
# =========================================================
@st.cache_resource
def load_ml_model(model_file, vector_file, dataset_name):
    with open(model_file, "rb") as f: model = pickle.load(f)
    with open(vector_file, "rb") as f: vec = pickle.load(f)
    return model, vec

# =========================================================
# Load DL model (local only)
# =========================================================
@st.cache_resource
def load_dl_model(model_file, tokenizer_file, dataset_name):
    model = load_model(model_file, compile=False)
    with open(tokenizer_file, "rb") as f: tok = pickle.load(f)
    return model, tok

# =========================================================
# Hide download button / toolbar
# =========================================================
st.markdown("""
    <style>
        [data-testid="stElementToolbar"] {display: none !important;}
        [data-testid="stDataFrameToolbar"] {display: none !important;}
    </style>
""", unsafe_allow_html=True)

# =========================================================
# Sidebar – model selection
# =========================================================
st.sidebar.title("Model Selection")
model_type = st.sidebar.radio("Select Model Type:", ["ML (Traditional)", "DL (BiLSTM)"])
if model_type == "ML (Traditional)":
    dataset = list(ml_datasets.keys())[0]
    cfg = ml_datasets[dataset]
    model, vectorizer = load_ml_model(cfg["model"], cfg["vectorizer"], dataset)
    is_ml = True
else:
    dataset = list(dl_datasets.keys())[0]
    cfg = dl_datasets[dataset]
    model, tokenizer = load_dl_model(cfg["model"], cfg["tokenizer"], dataset)
    is_ml = False

df, label_col = load_dataset(cfg, dataset)
text_col = cfg["text_col"]

# =========================================================
# Prediction function
# =========================================================
def predict(text):
    clean = clean_text(text)  # <-- clean input text
    if is_ml:
        pred = model.predict(vectorizer.transform([clean.lower()]))[0]
        return normalize_prediction(dataset, pred)
    else:
        seq = tokenizer.texts_to_sequences([clean])
        X = pad_sequences(seq, maxlen=MAX_LEN, padding="post", truncating="post")
        pred_prob = float(model.predict(X, verbose=0)[0][0])
        pred_class = 1 if pred_prob > 0.5 else 0
        label_map = DL_POSITIVE_LABEL.get(dataset, PRED_LABEL_MAP.get(dataset, {}))
        return label_map.get(pred_class, "Disinformation" if pred_class==1 else "True")

# =========================================================
# Dataset exploration
# =========================================================
st.subheader("Dataset Label Distribution")
df_vc = df[df[label_col].notna()]
st.dataframe(df_vc[label_col].value_counts().rename("Count"), width='stretch')

# =========================================================
# Text search / filter
# =========================================================
valid_labels = [l for l in df[label_col].unique() if str(l).strip().lower() not in ["","nan"]]
labels = ["All"] + sorted(valid_labels)
label_filter = st.radio("Filter by label:", labels, horizontal=True)
df_f = df if label_filter=="All" else df[df[label_col]==label_filter]
df_f = df_f[df_f[label_col].notna()]
df_f.reset_index(drop=True, inplace=True)

search_query = st.text_input("Search in text column:")
if search_query:
    df_f = df_f[df_f[text_col].str.contains(search_query, case=False, na=False)]
    df_f.reset_index(drop=True, inplace=True)

# =========================================================
# Dataset view / data editor
# =========================================================
df_view = df_f.groupby(label_col, group_keys=False).head(10) if label_filter=="All" else df_f.head(20)
df_view = df_view.reset_index(drop=True)

# ---------------------------------------------------------
# 1. Show dataset (read-only table)
# ---------------------------------------------------------


# ---------------------------------------------------------
# 2. Radio selection (TRUE SINGLE CHOICE)
# ---------------------------------------------------------
dataset_mode = (st.session_state.active_source == "dataset")
st.subheader("1. Select a record for prediction with Trained Dataset")

options = df_view.index.tolist()
selected_index = st.radio(
    "Choose row index:",
    df_view.index.tolist(),
    format_func=lambda x: df_view.loc[x, text_col][:180],
    key="dataset_radio",
    disabled=(st.session_state.active_source != "dataset")
)

# ONLY update if dataset is active
if st.session_state.active_source == "dataset":
    dataset_text = clean_text(df_view.loc[selected_index, text_col])

    if st.session_state.input_text != dataset_text:
        set_active_source("dataset", dataset_text)
    




# =========================================================
# File uploader
# =========================================================
import zipfile
import io
# =========================================================
# ZIP Folder Upload + Internal File Browser
# =========================================================

# =========================================================
# ZIP Folder Upload + Internal File Browser
# =========================================================

import zipfile
import io

ZIP_PATH = "news_data.zip"   # 👈 your file name

try:
    if os.path.exists(ZIP_PATH):
        zip_data = zipfile.ZipFile(ZIP_PATH)
    else:
        st.warning("ZIP file not found")
        zip_data = None

    file_list = []

    if zip_data is not None:

        # =====================================================
        # GET ONLY MAIN FOLDER + FIRST 5 SUBFOLDERS
        # =====================================================

        folder_structure = {}

        for f in zip_data.namelist():

            # remove invalid/system files
            if (
                "__MACOSX" in f
                or f.startswith(".")
                or os.path.basename(f).startswith("~$")
            ):
                continue

            parts = f.split("/")

            # Expect:
            # Main Folder/Sub Folder/file.pdf
            if len(parts) >= 4:

                main_folder = parts[1]
                sub_folder = parts[2]

                if main_folder not in folder_structure:
                    folder_structure[main_folder] = set()

                folder_structure[main_folder].add(sub_folder)

        # =====================================================
        # DISPLAY FOLDER STRUCTURE
        # =====================================================

        st.subheader("📂 ZIP Folder Structure")

        for main_folder, subfolders in folder_structure.items():

            st.markdown(f"### 📁 {main_folder}")

            for sf in sorted(list(subfolders))[:5]:
                st.markdown(f"- 📂 {sf}")

        # =====================================================
        # FILE LIST FOR SELECTION
        # =====================================================

        file_list = [
            f for f in zip_data.namelist()
            if (
                f.endswith((".txt", ".pdf", ".docx", ".csv"))
                and "__MACOSX" not in f
                and not f.startswith(".")
                and not os.path.basename(f).startswith("~$")
    )
]
        display_files = {}

        for f in file_list:

            parts = f.split("/")   
                    # news_data/Main Folder/Sub Folder/file.pdf
            if len(parts) >= 4:

                clean_name = f"{parts[1]} / {parts[2]} / {parts[-1]}"

            else:
                clean_name = f

            display_files[clean_name] = f 

    st.subheader("📁 2. Select a record for prediction with unseen data")

    
    selected_display = st.selectbox(
    "Select a file for prediction:",
    ["-- Select ZIP file --"] + list(display_files.keys()),
    key="zip_selector"
)
    # Get original ZIP path
    selected_file = (
        display_files[selected_display]
        if selected_display != "-- Select ZIP file --"
        else "-- Select ZIP file --"
)

    if os.path.basename(selected_file).startswith("~$"):
        st.warning("Skipping temporary system file")
        st.stop()

    if "last_zip_file" not in st.session_state:
        st.session_state.last_zip_file = None

    if (
        selected_file != "-- Select ZIP file --"
        and selected_file != st.session_state.last_zip_file
    ):

        st.session_state.active_source = "zip"
        st.session_state.last_zip_file = selected_file

        st.session_state.current_label = None
        st.session_state.summary = None
        st.session_state.prediction_done = False

        file_bytes = zip_data.read(selected_file)
        ext = selected_file.split(".")[-1].lower()

        text = ""

        if ext == "txt":
            text = file_bytes.decode("utf-8", errors="ignore")

        elif ext == "pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            text = "\n".join(
                p.extract_text() or ""
                for p in reader.pages
            )

        elif ext == "docx":
            doc = Document(io.BytesIO(file_bytes))
            text = "\n".join(
                p.text for p in doc.paragraphs
            )

        elif ext == "csv":
            text = file_bytes.decode(
                "utf-8",
                errors="ignore"
            )

        # ✅ UPDATE SESSION TEXT
        st.session_state.input_source = "zip"
        st.session_state.input_text = clean_text(text)

        st.success(f"Loaded: {selected_file}")

        def set_input(source, text):
            st.session_state.active_source = source
            st.session_state.input_text = clean_text(text)
            st.session_state.text_hash = None

except FileNotFoundError:
    st.error("news_data.zip not found. Please place it in the same folder as app.py")

# =========================================================
# Text area
# =========================================================
# =========================================================
# User Upload Option
# =========================================================
st.subheader("📤 3. Select and upload your own data")

uploaded_file = st.file_uploader(
    "Upload TXT, PDF, DOCX, or CSV",
    type=["txt", "pdf", "docx", "csv"]
)

if uploaded_file is not None:

    # Override previous folder-selected text
    if st.session_state.active_source != "upload":

        st.session_state.active_source = "upload"

        st.session_state.prediction_done = False
        st.session_state.current_label = None
        st.session_state.summary = None

    file_name = uploaded_file.name.lower()
    text = ""

    try:
        # TXT
        if file_name.endswith(".txt"):
            text = uploaded_file.read().decode(
                "utf-8",
                errors="ignore"
            )

        # PDF
        elif file_name.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            text = "\n".join(
                page.extract_text() or ""
                for page in reader.pages
            )

        # DOCX
        elif file_name.endswith(".docx"):
            doc = Document(uploaded_file)
            text = "\n".join(
                para.text for para in doc.paragraphs
            )

        # CSV
        elif file_name.endswith(".csv"):
            uploaded_df = pd.read_csv(uploaded_file)

            possible_cols = [
                "text",
                "content",
                "article",
                "news",
                "text_english"
            ]

            found_col = None

            for col in possible_cols:
                if col in uploaded_df.columns:
                    found_col = col
                    break

            if found_col:
                text = " ".join(
                    uploaded_df[found_col]
                    .astype(str)
                    .tolist()
                )
            else:
                st.error("No valid text column found in CSV")

        text = clean_text(text)

        if text.strip():
            st.session_state.input_text = text
            st.success(
                f"Uploaded file loaded: {uploaded_file.name}"
            )
            
            st.session_state.prediction_done = False         
        else:
            st.warning("No readable text found")

    except Exception as e:
        st.error(f"Error reading uploaded file: {e}")
# =========================================================
# Display selected or uploaded text (manual input disabled)
# =========================================================

#AI summary function
# =========================================================
# AI summary function (FIXED)
# =========================================================
@st.cache_data(show_spinner=False)
def generate_summary(keywords, label):
    if not keywords:
        return "No significant keywords were extracted."

    try:
        keyword_list = [k.strip() for k in keywords.split(",") if k.strip()]
        top_keywords = " ".join(keyword_list[:5])

        prompt = f"""
Keywords:
{top_keywords}

Label:
{label}
Display the label also with predected label {label}
Explain briefly in under 60 words.
"""

        # 🔥 SAFE GUARD
        response = gemini_model.generate_content(prompt)

        if response and hasattr(response, "text"):
            return response.text.strip()

        return "No response from AI model."

    except Exception as e:
        return f"AI error: {str(e)}"
# =========================================================
# Dynamic color Predict button
# =========================================================
button_color = "#4CAF50"  # default green
button_placeholder = st.empty()


## =========================================================
# Predict Button
# =========================================================
# =========================================================
# DISPLAY SELECTED / UPLOADED TEXT
# =========================================================
st.subheader("📝 Selected Text")

if st.session_state.input_text.strip() != "":

    st.text_area(
        "Selected text for prediction:",
        value=st.session_state.input_text,
        height=250,
        disabled=True
    )

else:

    st.info(
        "Select a dataset record, ZIP file, or upload your own file."
    )
# =========================================================
# Predict Button
# =========================================================




if st.button("Predict", key="predict_btn"):

    current_text = clean_text(st.session_state.input_text)
    current_source = st.session_state.active_source

    st.write(f"DEBUG SOURCE USED: {current_source}")
    st.write(f"DEBUG SOURCE USED FOR PREDICTION: {current_source}")
    st.write("DEBUG SOURCE:", st.session_state.active_source)
    st.write("DEBUG TEXT HASH:", get_hash(st.session_state.input_text))


      # 🔥 ensure fresh input
    if st.session_state.text_hash != get_hash(current_text):
        st.session_state.text_hash = get_hash(current_text)

    if current_text == "":
        st.warning("Please select or upload text first.")
    else:
        label = predict(current_text)

        # ============================
        # SHOW RESULT IMMEDIATELY
        # ============================
        label_color = (
            "#FF4B4B"
            if label.lower() in ["fake", "disinformation"]
            else "#4CAF50"
        )

        label_icon = (
            "🚨"
            if label.lower() in ["fake", "disinformation"]
            else "✅"
        )

        st.markdown(f"""
        <div style="
            background-color: {label_color};
            padding: 18px;
            border-radius: 12px;
            text-align: center;
            font-size: 28px;
            font-weight: bold;
            color: white;
            margin-top: 15px;
        ">
            {label_icon} Prediction: {label.upper()}
        </div>
        """, unsafe_allow_html=True)

        # ============================
        # VISUALIZATION
        # ============================
        st.subheader("📊 Important Keywords")

        if is_ml:

            X = vectorizer.transform([current_text.lower()])
            names = vectorizer.get_feature_names_out()

            if X.nnz > 0:

                df_plot = pd.DataFrame({
                    "Word": names[X.indices],
                    "Importance": X.data
                }).sort_values(
                    "Importance",
                    ascending=False
                ).head(20)

                chart = alt.Chart(df_plot).mark_bar().encode(
                    x="Importance",
                    y=alt.Y("Word", sort='-x')
                )

                st.altair_chart(
                    chart,
                    use_container_width=True
                )

                st.session_state.top_keywords = ", ".join(
                    df_plot["Word"].tolist()
                )

            else:
                st.info("No important keywords found.")
                st.session_state.top_keywords = ""

        else:

            wc = WordCloud(
                width=800,
                height=400,
                background_color="white",
                max_words=20
            ).generate(current_text)

            fig, ax = plt.subplots()

            ax.imshow(wc)

            ax.axis("off")

            st.pyplot(fig)

            plt.close(fig)

            st.session_state.top_keywords = ", ".join(
                list(wc.words_.keys())[:20]
            )

        # ============================
        # AI EXPLANATION LAST
        # ============================
        st.subheader("🧠 AI Explanation")

        with st.spinner("Generating explanation..."):

            summary = generate_summary(
                st.session_state.top_keywords,
                label
            )

            st.session_state.summary = summary

        st.write(st.session_state.summary)
import gspread
from google.oauth2.service_account import Credentials
import datetime

SHEET_ID = "1UklQD2URWeB9l6X4hTnDc9YRcNJA0d7BeOx7wXTz280"
WORKSHEET_NAME = "X_Feedback_Dashboard"

@st.cache_resource
def connect_gsheet():

    creds_dict = dict(st.secrets["gcp_service_account"])

    creds = Credentials.from_service_account_info(
        creds_dict,
        scopes=[
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"
        ]
    )

    client = gspread.authorize(creds)

    spreadsheet = client.open_by_key(SHEET_ID)

    return spreadsheet.worksheet(WORKSHEET_NAME)

try:
    sheet = connect_gsheet()

except Exception as e:
    sheet = None
    st.error(f"Google Sheet connection failed: {e}")

# =========================================================
st.markdown("---")
st.header("🧠 Post-Prediction Questionnaire")
st.subheader("System Evaluation (Likert Scale)")

if "questionnaire_submitted" not in st.session_state:
    st.session_state.questionnaire_submitted = False

with st.form("feedback_form"):

    st.subheader("User Profile")

    age = st.selectbox("Age Group", ["18–24","25–34","35–44","45–54","55+"])
    gender = st.selectbox("Gender", ["Male","Female","Other","Prefer not to say"])
    education = st.selectbox(
        "Education Level",
        ["High School","Undergraduate","Postgraduate","PhD","Other"]
    )

    field = st.selectbox(
        "Field of Study",
        ["AI/CS","Social Sciences","Journalism","Business","Other"]
    )

    ai_level = st.selectbox(
        "AI Knowledge",
        ["None","Basic","Intermediate","Advanced"]
    )

    news_freq = st.selectbox(
        "News Consumption Frequency",
        ["Rarely","Weekly","Daily","Multiple times/day"]
    )

    st.subheader("System Evaluation (1–5)")

    st.markdown("""
**1 = Strongly Disagree**  
**2 = Disagree**  
**3 = Neutral**  
**4 = Agree**  
**5 = Strongly Agree**
""")

    # -------------------------------
    # SUS-Inspired Usability Questions
    # -------------------------------
    st.markdown("### Usability (SUS-Inspired)")
    scale_hint = "(1 = Strongly Disagree → 5 = Strongly Agree)"

    q1 = st.slider(
        f"I think that I would like to use this system frequently {scale_hint}.",
        1, 5, 3
    )

    q2 = st.slider(
        f"I found the system unnecessarily complex {scale_hint}.",
        1, 5, 3
    )

    q3 = st.slider(
        f"I thought the system was easy to use {scale_hint}.",
        1, 5, 3
    )

    q4 = st.slider(
        f"I think that I would need the support of a technical person to use this system {scale_hint}.",
        1, 5, 3
    )

    q5 = st.slider(
        f"I found the various functions in this system were well integrated {scale_hint}.",
        1, 5, 3
    )

    q6 = st.slider(
        f"I thought there was too much inconsistency in this system {scale_hint}.",
        1, 5, 3
    )
    q7 = st.slider(
        f"I would imagine that most people would learn to use this system very quickly {scale_hint}.",
        1, 5, 3
    )

    q8 = st.slider(
        f"I found the system very cumbersome (Difficult) to use {scale_hint}.",
        1, 5, 3
    )
    q9 = st.slider(
        f"I felt very confident using the system {scale_hint}.",
        1, 5, 3
    )
    q10 = st.slider(
        f"I needed to learn a lot of things before I could get going with this system {scale_hint}.",
        1, 5, 3
    )

    # -------------------------------
    # Explainability
    # -------------------------------
    st.markdown("### Explainability")

    q11 = st.slider(
        f"The model predictions were clearly presented {scale_hint}.",
        1, 5, 3
    )

    q12 = st.slider(
        f"The visualisations and summary helped me understand the prediction results {scale_hint}.",
        1, 5, 3
    )
    # -------------------------------
    # Trust
    # -------------------------------
    st.markdown("### Trust")

    q13 = st.slider(
        f"I trust the system’s predictions {scale_hint}.",
        1, 5, 3
    )

    # -------------------------------
    # Usefulness
    # -------------------------------
    st.markdown("### Usefulness")

    q14 = st.slider(
        f"This system is useful for identifying fake news/disinformation {scale_hint}.",
        1, 5, 3
    )

    q15 = st.slider(
        f"The dashboard helped me make better decisions about whether news content was trustworthy {scale_hint}.",
        1, 5, 3
    )

    # -------------------------------
    # Adoption
    # -------------------------------
    st.markdown("### Future Adoption")

    q16 = st.slider(
        f"I would use this system again in the future {scale_hint}.",
        1, 5, 3
    )

    comments = st.text_area(
        "What improvements would you suggest for this dashboard?"
    )

    submit = st.form_submit_button("Submit Feedback")

# =========================================================
# SAVE TO GOOGLE SHEETS
# =========================================================
if submit:

    try:
        if sheet is not None:
            sheet.append_row([
                str(datetime.datetime.now()),
                model_type,
                dataset,
                age,
                gender,
                education,
                field,
                ai_level,
                news_freq,

                # SUS-inspired usability
                q1,  # FrequentUse
                q2,  # Complex
                q3,  # EasyUse
                q4,  # TechSupport
                q5,  # WellIntegrated
                q6,  # Inconsistency
                q7,  # QuickLearning
                q8,  # Cumbersome (Difficult)
                q9,  # Confidence
                q10,  # LearnBeforeUse

                # Explainability
                q11,  # Prediction_Clarity
                q12,  # Visualization_Helpfulness

                # Trust
                q13,  # Trust

                # Usefulness
                q14,  # Usefulness
                q15, # Decision_Support

                # Adoption
                q16, # Reuse_Intention

                comments
            ])

            st.success("✅ Feedback saved")

        else:
            st.warning("⚠ Google Sheet not connected. Data not saved.")

        st.session_state.questionnaire_submitted = True

    except Exception as e:
        st.error(f"Error saving feedback: {e}")